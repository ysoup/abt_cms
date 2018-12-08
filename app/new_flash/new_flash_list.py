# encoding=utf-8
from flask_login import login_required
from flask import request, render_template, jsonify, make_response,flash, abort, url_for, redirect, session, Flask, g, current_app
from . import new_flash
from app.models import AccountManage, ArticleManage, ArticleUploadManage, InformationPlatform, NewInformationCategory, \
    ArticleUploadDetails
from cms_server import db, redis_store
import time, copy
from common import push_service
import json, datetime
import os
import xlrd
# utils相关
from ..utils.filter_char import filter_char
from ..utils.sphinxapi3_query_keyword import query_keyword
from ..utils.Pagination import Pagination
from werkzeug.utils import secure_filename
import docx
import mammoth
import mammoth.transforms
import os
from docx import Document
import zipfile
import requests
import filetype

styleMap = """

p[style-name='Title'] => h1.hide

p[style-name='Subhead 1'] => h3

p[style-name='List Bullet'] => ul.first > li:fresh

p[style-name='List Bullet 2'] => ul.second > li:fresh

p[style-name='Hyperlink']=>a.link

"""
guidUrl = "https://my.phrplus.com/REST/guid"

imgPath = "/data/imgs"


# 获取回调
@new_flash.route('/redirect_uri', methods=['GET'])
def redirect_uri():
    try:
        return render_template('account/redirect_uri.html')
    except Exception as e:
        current_app.logger.error(e)
        return render_template("404.html")


# 删除资讯
@new_flash.route("/delete_account", methods=['GET', 'POST'])
@login_required
def delete_account():
    if request.method == "POST":
        try:
            info = AccountManage.query.filter_by(id=request.form['id']).first()
            db.session.delete(info)
            db.session.commit()
            return jsonify({'success': 'ok'})
        except Exception as e:
            current_app.logger.error(e)
            db.session.rollback()
        return jsonify({'failed': 'ok'})


# 账户列表页
@new_flash.route('/account_manage_list', methods=['GET'])
@login_required
def account_manage_list():
    try:
        page = request.args.get('page', 1, type=int)
        pagination = AccountManage.query.order_by(AccountManage.create_time.desc()).paginate(page, per_page=10,
                                                                                             error_out=False)
        data = pagination.items
        # 账户类型
        types_ls = [
            {0: "3级图文原创"},
            {1: "2级图文原创"},
            {2: "3级非图文原创"},
            {3: "2级非图文原创"}
        ]
        # 账户平台
        platform_rows = InformationPlatform.query
        platform_ls = []
        if platform_rows:
            for x in platform_rows:
                dic = {}
                dic["id"] = x.id
                dic["platform_name"] = x.platform_name
                platform_ls.append(dic)

        # 账户分类
        platform_category=NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dic = {}
                dic["id"] = x.id
                dic["category_name"] = x.category_name
                category_ls.append(dic)

        return render_template('account/account_manage_list.html', data=data,
                               types_ls=types_ls,
                               platform_ls=platform_ls,
                               category_ls=category_ls,
                               pagination=pagination)
    except Exception as e:
        current_app.logger.error(e)
        return render_template("404.html")


# 新增账户信息
@new_flash.route('/add_account_info', methods=['GET', "POST"])
def add_account_info():
    if request.method == "GET":
        # 账户平台
        platform_rows = InformationPlatform.query
        platform_ls = []
        if platform_rows:
            for x in platform_rows:
                dic = {}
                dic["id"] = x.id
                dic["platform_name"] = x.platform_name
                platform_ls.append(dic)

        # 账户分类
        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dic = {}
                dic["id"] = x.id
                dic["category_name"] = x.category_name
                category_ls.append(dic)

        return render_template('account/add_account_info.html',platform_ls=platform_ls,
                               category_ls=category_ls)
    elif request.method == "POST":
        try:
            account_name = request.form.get("account_name")
            account_password = request.form.get("account_password")
            account_rank = request.form.get("account_rank")
            account_type = request.form.get("account_type")
            total_read_num = request.form.get("total_read_num")
            total_play_num = request.form.get("total_play_num")
            total_subscribe_num = request.form.get("total_subscribe_num")
            account_index = request.form.get("account_index")
            # account_article_num = request.form.get("account_article_num")
            nick_name = request.form.get("nick_name")
            credit_score = request.form.get("credit_score")
            platform_type = request.form.get("platform_type")
            category_type = request.form.get("category_type")

            info = AccountManage(
                account_name=account_name,
                account_password=account_password,
                account_rank=account_rank,
                # account_article_num=account_article_num,
                account_type=account_type,
                total_read_num=total_read_num,
                total_play_num=total_play_num,
                total_subscribe_num=total_subscribe_num,
                account_index=account_index,
                nick_name=nick_name,
                credit_score=credit_score,
                platform_type=platform_type,
                category_type=category_type,
                is_delete=0
            )
            db.session.add(info)
            db.session.commit()
            return jsonify({'success': 'ok'})
        except Exception as e:
            current_app.logger.error(e)
            return jsonify({'failed': 'ok'})


# 编辑账户信息
@new_flash.route("/modify_account_info", methods=['GET', 'POST'])
def modify_account_info():
    if request.method == "GET":
        id = request.args.get('id', type=int)
        info = AccountManage.query.filter_by(id=id).first()
        dics = {}
        if info:
            dics["id"] = id
            dics["account_name"] = info.account_name
            dics["account_password"] = info.account_password
            dics["account_rank"] = info.account_rank
            dics["account_type"] = info.account_type
            dics["total_read_num"] = info.total_read_num
            dics["total_play_num"] = info.total_play_num
            dics["total_subscribe_num"] = info.total_subscribe_num
            dics["account_index"] = info.account_index
            dics["account_article_num"] = info.account_article_num
            dics["nick_name"] = info.nick_name
            dics["credit_score"] = info.credit_score
            dics["platform_type"] = info.platform_type
            dics["category_type"] = info.category_type

        # 账户平台
        platform_rows = InformationPlatform.query
        platform_ls = []
        if platform_rows:
            for x in platform_rows:
                dic = {}
                dic["id"] = x.id
                dic["platform_name"] = x.platform_name
                platform_ls.append(dic)

        # 账户分类
        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dic = {}
                dic["id"] = x.id
                dic["category_name"] = x.category_name
                category_ls.append(dic)
        return render_template('account/modify_account_manage.html', data=dics, platform_ls=platform_ls,
                               category_ls=category_ls)
    elif request.method == "POST":
        try:
            id = request.form.get('id')
            account_name = request.form.get('account_name')
            account_password = request.form.get('account_password')
            account_rank = request.form.get('account_rank')
            # account_article_num = request.form.get('account_article_num')
            account_type = request.form.get('account_type')
            total_read_num = request.form.get('total_read_num')
            total_play_num = request.form.get('total_play_num')
            total_subscribe_num = request.form.get('total_subscribe_num')
            account_index = request.form.get('account_index')

            nick_name = request.form.get('nick_name')
            credit_score = request.form.get('credit_score')
            platform_type = request.form.get('platform_type')
            category_type = request.form.get('category_type')
            info = AccountManage.query.filter_by(id=id).first()
            if info:
                info.account_name = account_name
                info.account_password = account_password
                info.account_rank = account_rank
                # info.account_article_num = account_article_num
                info.account_type = account_type
                info.total_read_num = total_read_num
                info.total_play_num = total_play_num
                info.total_subscribe_num = total_subscribe_num
                info.account_index = account_index
                info.nick_name = nick_name
                info.credit_score = credit_score
                info.platform_type = platform_type
                info.category_type = category_type
                db.session.add(info)
                db.session.commit()
            return jsonify({"success": "ok"})
        except Exception as e:
            current_app.logger.error(e)
            return jsonify({'failed': '修改失败'})


# 文件上传
@new_flash.route('/file_upload', methods=['POST'])
@login_required
def file_upload():
    try:
        file_dict = request.files["file_data"]
        account_type = request.form["account_type"]
        platform_type = request.form["platform_type"]
        category_type = request.form["category_type"]
        filename = secure_filename(file_dict.filename)
        file_dict.save(os.path.join("./", filename))
        xls_file = xlrd.open_workbook(filename)
        xls_sheet = xls_file.sheets()[0]
        nrows = xls_sheet.nrows
        ncols = xls_sheet.ncols
        for i in range(1, nrows):
            info = []
            for j in range(0, ncols):
                var = xls_sheet.cell(i, j).value
                info.append(var)
            # 数据库持久化
            info = AccountManage(
                account_name=info[0],
                nick_name=info[1],
                account_password=info[2],
                account_rank=info[3],
                total_read_num=info[4],
                total_play_num=info[5],
                total_subscribe_num=info[6],
                account_index=info[7],
                credit_score=info[8],
                account_type=account_type,
                platform_type=platform_type,
                category_type=category_type)
            db.session.add(info)
            db.session.commit()
        return jsonify({"success": "ok"})
    except Exception as e:
        current_app.logger.error(e)
        return jsonify({"success": "failed"})


# 数据统计列表
@new_flash.route('/account_count_list', methods=['GET'])
@login_required
def account_count_list():
    try:
        page = request.args.get('page', 1, type=int)
        pagination = AccountManage.query.order_by(AccountManage.create_time.desc()).paginate(page, per_page=10,
                                                                                             error_out=False)
        data = pagination.items
        return render_template('account/account_count_list.html', data=data, pagination=pagination)
    except Exception as e:
        current_app.logger.error(e)
        return render_template("404.html")


# 文章列表
@new_flash.route('/article_list', methods=['GET'])
@login_required
def article_list():
    try:
        page = request.args.get('page', 1, type=int)
        if request.args.__contains__("category_type"):
            category_type = request.args.get("category_type", type=int)
            is_send = request.args.get("is_send", type=int)
            control_status = request.args.get("control_status", type=int)
            if category_type != 888 and is_send != 888 and control_status != 888:
                pagination = ArticleManage.query.filter_by(category_type=category_type,
                                                           is_send=is_send,
                                                           control_status=control_status).\
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type == 888 and is_send != 888 and control_status != 888:
                pagination = ArticleManage.query.filter_by(
                                                           is_send=is_send,
                                                           control_status=control_status). \
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type == 888 and is_send == 888 and control_status != 888:
                pagination = ArticleManage.query.filter_by(
                    control_status=control_status). \
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type != 888 and is_send == 888 and control_status != 888:
                pagination = ArticleManage.query.filter_by(category_type=category_type,
                                                           control_status=control_status). \
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type != 888 and is_send != 888 and control_status == 888:
                pagination = ArticleManage.query.filter_by(category_type=category_type,
                                                           is_send=is_send
                                                           ). \
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type != 888 and is_send == 888 and control_status == 888:
                pagination = ArticleManage.query.filter_by(category_type=category_type). \
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type == 888 and is_send != 888 and control_status == 888:
                pagination = ArticleManage.query.filter_by(is_send=is_send).\
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            elif category_type == 888 and is_send == 888 and control_status != 888:
                pagination = ArticleManage.query.filter_by(control_status=control_status). \
                    order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30, error_out=False)
            else:
                pagination = ArticleManage.query.order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30,
                                                                                                     error_out=False)
        else:
            pagination = ArticleManage.query.order_by(ArticleManage.create_time.desc()).paginate(page, per_page=30,
                                                                                                 error_out=False)
        data = pagination.items
        # 账户分类
        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dic = {}
                dic["id"] = x.id
                dic["category_name"] = x.category_name
                category_ls.append(dic)

        # 账户平台
        platform_rows = InformationPlatform.query
        platform_ls = []
        if platform_rows:
            for x in platform_rows:
                dict = {}
                dict["id"] = x.id
                dict["platform_name"] = x.platform_name
                platform_ls.append(dict)

        # 账户列表
        account_rows = AccountManage.query
        acclout_ls_0 = {
            "type_name": "3级图文原创",
            "type_id": 0,
            "ls": []
        }
        acclout_ls_1 = {
            "type_name": "2级图文原创",
            "type_id": 0,
            "ls": []
        }
        acclout_ls_2 = {
            "type_name": "3级非图文原创",
            "type_id": 0,
            "ls": []
        }
        acclout_ls_3 = {
            "type_name": "2级非图文原创",
            "type_id": 0,
            "ls": []
        }
        dat = datetime.date.today()
        if account_rows:
            for x in account_rows:
                # 判断该账户发送文章次数
                workorders = ArticleUploadDetails.query.filter(
                    db.cast(ArticleUploadDetails.create_time, db.DATE) == dat).filter(ArticleUploadDetails.account_id==x.id).all()
                if len(workorders) ==2: continue
                dics = {}
                dics["id"] = x.id
                dics["account_name"] = x.account_name
                if x.account_type == 0:
                    acclout_ls_0["ls"].append(dics)
                elif x.account_type == 1:
                    acclout_ls_1["ls"].append(dics)
                elif x.account_type == 2:
                    acclout_ls_2["ls"].append(dics)
                else:
                    acclout_ls_3["ls"].append(dics)
        acclout_ls = [acclout_ls_0, acclout_ls_1, acclout_ls_2, acclout_ls_3]
        return render_template('article/article_list.html', data=data, category_ls=category_ls, platform_ls=platform_ls,
                               acclout_ls=acclout_ls, pagination=pagination)
    except Exception as e:
        current_app.logger.error(e)
        return render_template("404.html")


# 文章上传
@new_flash.route('/article_file_upload', methods=['POST'])
@login_required
def article_file_upload():
    try:
        file_dict = request.files["file_data"]
        account_type = request.form["account_type"]
        category_type = request.form["category_type"]
        filename = secure_filename(file_dict.filename)
        file_dict.save(os.path.join("./", filename))
        kind = filetype.guess(filename)
        print(kind.extension)
        article = parseFile(filename)
        imgNameArr = extractImage(filename)

        fileName = os.path.basename(filename)

        contentArr = article["Content"].split("<img")

        for idx, section in enumerate(contentArr):

            for info in imgNameArr:
                if idx is info["index"]:
                    contentArr[idx] = section + "<img alt='" + info["fileName"] + "' data-fileName='" + info[
                        "fileName"] + "'"
        article["Content"] = ''.join(contentArr)

        # 数据库持久化
        info = ArticleManage(
            article_title=article["Title"],
            article_cover="",
            article_content=article["Content"],
            article_type=account_type,
            category_type=category_type,
            is_send=1
        )
        db.session.add(info)
        db.session.commit()
        return jsonify({"success": "ok"})
    except Exception as e:
        current_app.logger.error(e)
        return jsonify({"success": "failed"})


def extractImage(f):
    kind = filetype.guess(filename)
    ziped = zipfile.ZipFile(f)
    allFiles = ziped.namelist()
    imgs = filter(lambda x: x.startswith('word/media/'), allFiles)
    imgNameArr = []
    for img in imgs:
        res = requests.post(guidUrl)
        if res.status_code is 200:
            guid = res.text
            data = ziped.read(img, bytes(imgPath, encoding="utf-8"))
            idxStr = os.path.basename(img).split(".")[0][-1:]
            imgDict = {}
            imgDict["index"] = int(idxStr)-1
            imgDict["fileName"] = guid+".jpg"
            imgNameArr.append(imgDict)
            targetPath = os.path.join(imgPath, guid+".jpg")
            target = open(targetPath, "wb")
            target.write(data)
            target.close()
    ziped.close()

    return imgNameArr


def parseFile(f):
    document = Document(f)
    article = {"Title": document.paragraphs[0].text, "Content":""}
    with open(f, "rb") as docFile:
        html = mammoth.convert_to_html(docFile, style_map=styleMap,
                                       convert_image=mammoth.images.img_element(convert_image))
    if not article["Title"]:
        for para in document.paragraphs:
            if para.style.name == 'Title':
                if para.text:
                    article["Title"] = para.text

    article["Content"] = html.value
    return article


def convert_image(image):
    return {"src": ""}


# 新增文章信息
@new_flash.route('/add_article_info', methods=['GET', "POST"])
def add_article_info():
    if request.method == "GET":
        # 账户分类
        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dic = {}
                dic["id"] = x.id
                dic["category_name"] = x.category_name
                category_ls.append(dic)
        return render_template('article/add_article_info.html', category_ls=category_ls)
    elif request.method == "POST":
        try:
            article_content = request.form.get("article_content")
            article_title = request.form.get("article_title")
            article_type = request.form.get("article_type")
            category_type = request.form.get("category_type")

            info = ArticleManage(
                article_content=article_content,
                article_title=article_title,
                article_type=article_type,
                is_send=0,
                category_type=category_type
            )
            db.session.add(info)
            db.session.commit()
            return jsonify({'success': 'ok'})
        except Exception as e:
            current_app.logger.error(e)
            return jsonify({'failed': 'ok'})


# 编辑文章信息
@new_flash.route("/modify_article_info", methods=['GET', 'POST'])
def modify_article_info():
    if request.method == "GET":
        id = request.args.get('id', type=int)
        info = ArticleManage.query.filter_by(id=id).first()
        dic = {}
        if info:
            dic["id"] = id
            dic["article_title"] = info.article_title
            dic["article_content"] = info.article_content
            dic["article_type"] = info.article_type
            dic["category_type"] = info.category_type

        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dics = {}
                dics["id"] = x.id
                dics["category_name"] = x.category_name
                category_ls.append(dics)
        return render_template('article/modify_article_manage.html', data=dic, category_ls=category_ls)
    elif request.method == "POST":
        try:
            id = request.form.get('id')
            article_title = request.form.get('article_title')
            article_content = request.form.get('article_content')
            category_type = request.form.get('category_type')
            article_type = request.form.get('article_type')
            info = ArticleManage.query.filter_by(id=id).first()
            if info:
                info.article_title = article_title
                info.article_content = article_content
                info.article_type = article_type
                info.category_type = category_type
                info.is_send = 1
                info.control_status = 1
                db.session.add(info)
                db.session.commit()
            return jsonify({"success": "ok"})
        except Exception as e:
            current_app.logger.error(e)
            return jsonify({'failed': '修改失败'})


# 文章上传设置列表
@new_flash.route('/article_upload_set', methods=['GET'])
@login_required                                                                                                 
def article_upload_set():
    try:
        page = request.args.get('page', 1, type=int)
        pagination = ArticleUploadManage.query.order_by(ArticleUploadManage.create_time.desc()).paginate(page,
                                                                                                         per_page=10,
                                                                                                         error_out=False)
        data = pagination.items

        # 账户类型
        types_ls = [
            {0: "3级图文原创"},
            {1: "2级图文原创"},
            {2: "3级非图文原创"},
            {3: "2级非图文原创"}
        ]

        category_ls = []
        platform_category = NewInformationCategory.query
        if platform_category:
            for x in platform_category:
                dics = {}
                dics["id"] = x.id
                dics["category_name"] = x.category_name
                category_ls.append(dics)

        # 账户平台
        platform_rows = InformationPlatform.query
        platform_ls = []
        if platform_rows:
            for x in platform_rows:
                dict = {}
                dict["id"] = x.id
                dict["platform_name"] = x.platform_name
                platform_ls.append(dict)

        return render_template('article/article_set_list.html', data=data, category_ls=category_ls,
                               platform_ls=platform_ls,
                               pagination=pagination)
    except Exception as e:
        current_app.logger.error(e)
        return render_template("404.html")


# 新增文章上传设置
@new_flash.route('/add_article_upload_set', methods=['GET', "POST"])
def add_article_upload_set():
    if request.method == "GET":
        info = AccountManage.query.order_by(AccountManage.create_time.desc())
        account_ls = []
        for x in info:
            dic = {}
            dic["id"] = x.id
            dic["account_name"] = x.account_name
            account_ls.append(dic)
        # 账户类型
        types_ls = [
            {0: "3级图文原创"},
            {1: "2级图文原创"},
            {2: "3级非图文原创"},
            {3: "2级非图文原创"}
        ]
        # 账户平台
        platform_rows = InformationPlatform.query
        platform_ls = []
        if platform_rows:
            for x in platform_rows:
                dic = {}
                dic["id"] = x.id
                dic["platform_name"] = x.platform_name
                platform_ls.append(dic)

        # 账户分类
        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dic = {}
                dic["id"] = x.id
                dic["category_name"] = x.category_name
                category_ls.append(dic)
        return render_template('article/add_article_upload_set.html', platform_ls=platform_ls,
                               data=account_ls, category_ls=category_ls)
    elif request.method == "POST":
        try:
            account_name = request.form.get("account_name")
            article_type = request.form.get("article_type")
            send_type = request.form.get("send_type")
            category_type = request.form.get("category_type")
            platform_name = request.form.get("platform_name")

            info = ArticleUploadManage(
                account_name=account_name,
                article_type=article_type,
                category_type=category_type,
                platform_type=platform_name,
                send_type=send_type
            )
            db.session.add(info)
            db.session.commit()
            return jsonify({'success': 'ok'})
        except Exception as e:
            current_app.logger.error(e)
            return jsonify({'failed': 'ok'})


# 编辑文章信息
@new_flash.route("/modify_article_upload_set", methods=['GET', 'POST'])
def modify_article_upload_set():
    if request.method == "GET":
        id = request.args.get('id', type=int)
        info = ArticleManage.query.filter_by(id=id).first()
        dic = {}
        if info:
            dic["id"] = id
            dic["article_title"] = info.article_title
            dic["article_content"] = info.article_content
            dic["article_type"] = info.article_type

        # 账户分类
        platform_category = NewInformationCategory.query
        category_ls = []
        if platform_category:
            for x in platform_category:
                dics = {}
                dics["id"] = x.id
                dics["category_name"] = x.category_name
                category_ls.append(dics)
        return render_template('article/modify_article_upload_set.html', data=dic, category_ls=category_ls)
    elif request.method == "POST":
        try:
            id = request.form.get('id')
            article_title = request.form.get('article_title')
            article_content = request.form.get('article_content')
            article_type = request.form.get('article_type')
            category_type = request.form.get('category_type')
            info = ArticleManage.query.filter_by(id=id).first()
            if info:
                info.article_title = article_title
                info.article_content = article_content
                info.article_type = article_type
                info.category_type = category_type
                db.session.add(info)
                db.session.commit()
            return jsonify({"success": "ok"})
        except Exception as e:
            current_app.logger.error(e)
            return jsonify({'failed': '修改失败'})


# 编辑文章信息
@new_flash.route("/set_category", methods=['POST'])
def set_category():
    try:
        category_type = request.form.get('category_type')
        info_id_list = request.form.get('info_id_list')
        id_list = info_id_list.split(",")
        for x in id_list:
            info = ArticleManage.query.filter_by(id=int(x)).first()
            info.category_type = category_type
        db.session.commit()
        return jsonify({"success": "ok"})
    except Exception as e:
        current_app.logger.error(e)
        return jsonify({'failed': '修改失败'})


# 编辑文章信息
@new_flash.route("/set_type", methods=['POST'])
def set_type():
    try:
        info_id_list = request.form.get('info_id_list')
        article_title = request.form.get('article_title')
        article_content = request.form.get('article_content')
        article_type = request.form.get('article_type')
        info = ArticleManage.query.filter_by(id=id).first()
        if info:
            info.article_title = article_title
            info.article_content = article_content
            info.article_type = article_type
            db.session.add(info)
            db.session.commit()
        return jsonify({"success": "ok"})
    except Exception as e:
        current_app.logger.error(e)
        return jsonify({'failed': '修改失败'})


# 编辑文章信息
@new_flash.route("/all_article_upload", methods=['POST'])
def all_article_upload():
    try:
        info_id_list = request.form.get('info_id_list')
        platform_type = request.form.get('platform_type')
        account_type = request.form.get('account_type')
        account_name = request.form.get('account_name')
        category_type = request.form.get('category_type')
        info = ArticleUploadManage(
            send_ids=info_id_list,
            article_type=account_type,
            category_type=category_type,
            platform_type=platform_type,
            account_name=account_name,
            send_type=0
        )
        db.session.add(info)
        db.session.commit()
        return jsonify({"success": "ok"})
    except Exception as e:
        current_app.logger.error(e)
        return jsonify({'failed': '修改失败'})


# 删除资讯
@new_flash.route("/delete_information", methods=['GET', 'POST'])
@login_required
def delete_information():
    if request.method == "POST":
        try:
            info = ArticleManage.query.filter_by(id=request.form['id']).first()
            db.session.delete(info)
            db.session.commit()
            return jsonify({'success': 'ok'})
        except Exception as e:
            current_app.logger.error(e)
            db.session.rollback()
        return jsonify({'failed': 'ok'})



